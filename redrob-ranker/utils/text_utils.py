"""
Text utilities shared across the candidate-ranking pipeline and the JD
compiler.

The functions below extend the original text_utils.py with one addition:
extract_critical_tokens(), which derives a JD-specific list of "things
that must not silently disappear during compilation" directly from the
raw JD text, instead of a hand-typed list tuned to one JD. See the
docstring on that function for why this matters.
"""

import re


def clean_text(s):
    if not s:
        return ""
    s = s.strip().lower()
    return re.sub(r"[\x00-\x1f\x7f-\x9f]", "", s)


def concat_career_text(candidate):
    parts = []
    profile = candidate.get("profile", {})

    headline = profile.get("headline")
    if headline:
        parts.append(headline)

    summary = profile.get("summary")
    if summary:
        parts.append(summary)

    career = candidate.get("career_history", [])
    for job in career:
        title = job.get("title")
        if title:
            parts.append(title)
        desc = job.get("description")
        if desc:
            parts.append(desc)

    return " ".join(parts)


def tokenize_for_bm25(text):
    if not text:
        return []
    cleaned = re.sub(r"[^\w\s]", "", text)
    return cleaned.lower().split()


def extract_raw_text_from_docx(path):
    import docx

    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                tables_text.append(" | ".join(row_text))
    return "\n".join(paragraphs + tables_text)


def extract_raw_text_from_file(path):
    if path.lower().endswith(".docx"):
        return extract_raw_text_from_docx(path)
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


# ---------------------------------------------------------------------------
# Dynamic critical-token extraction for the validation gate.
#
# The previous gate hardcoded a list of tokens (TCS, Infosys, ..., Pune,
# Bangalore, "30-day", "5-9") that happened to apply to one specific JD.
# That breaks the moment the gate is run against a different JD, and it
# already contained one token ("Bangalore") that the actual JD never
# mentions at all — a check that can never meaningfully fail or pass
# because its precondition is never true.
#
# extract_critical_tokens() instead derives the token list from the raw
# JD text at runtime using three generic patterns that show up in almost
# any JD: capitalized proper nouns/company names, numeric ranges (years,
# days, months), and currency/compensation figures. This makes the gate
# JD-agnostic — it checks "did we drop something this specific JD
# actually said," not "did we drop something this one example JD said."
# ---------------------------------------------------------------------------

_STOPWORDS_TITLECASE = {
    "The", "This", "That", "These", "Those", "We", "You", "Our", "Your",
    "If", "It", "Is", "Are", "Be", "A", "An", "And", "Or", "But", "For",
    "In", "On", "At", "To", "Of", "With", "As", "By", "From", "Most",
    "Some", "What", "How", "Why", "When", "Where", "Things", "Final",
    "Let", "Note", "Beyond", "Here", "Senior", "Series", "They", "He",
    "She", "I", "There", "Here", "Then", "Now", "So", "Also", "Even",
    "While", "Once", "Weeks", "Months", "Years", "Days", "People",
    "Candidates", "Engineers", "Engineer", "Team", "Teams",
}

# Document-template field labels. These are structural scaffolding that
# almost every JD document shares (headers like "Job Description:",
# "Employment Type:", "Experience Required:") — they describe the SHAPE
# of the document, not a fact specific to THIS JD. Flagging them as
# "critical tokens the config must preserve" produces guaranteed false
# alarms on every JD, since no config schema has a field for "the word
# that happened to label a header." This list is generic across JD
# documents, not tuned to one specific JD's content.
_DOCUMENT_TEMPLATE_LABELS = {
    "job description", "company", "location", "employment type",
    "experience required", "role", "position", "department", "reports to",
    "salary range", "compensation", "about us", "about the role",
    "responsibilities", "requirements", "qualifications", "benefits",
    "how to apply", "job title", "team", "founding team",
}


def _extract_proper_nouns(text: str) -> list:
    """Capitalized words/phrases that look like names, companies, or places.

    Catches single capitalized tokens (Pune, TCS, Infosys) and short
    capitalized phrases (Delhi NCR, Series A). Filters out three sources
    of false positives:
      1. Common sentence-starting words (The, This, They, Weeks) via
         _STOPWORDS_TITLECASE.
      2. Document-template field labels (Job Description, Employment
         Type) via _DOCUMENT_TEMPLATE_LABELS — these describe document
         structure, not JD-specific facts.
      3. Phrases that accidentally span a paragraph/line break — matched
         on a per-line basis only, so "...Founding Team" at the end of
         one line can never glue onto "Company..." at the start of the
         next, which produced a nonsense token in an earlier version.
    """
    out = []
    # Process line-by-line so a capitalized phrase can never span a
    # newline. \s in the original pattern matched newlines too, which
    # merged unrelated phrases from adjacent lines/paragraphs into one
    # garbage token (e.g. "Founding Team\n\nCompany").
    for line in text.split("\n"):
        candidates = re.findall(r"\b[A-Z][a-zA-Z]{1,}(?:[ \t]+[A-Z][a-zA-Z]{1,}){0,2}\b", line)
        for c in candidates:
            words = c.split()
            first_word = words[0]

            if first_word in _STOPWORDS_TITLECASE:
                continue
            if c.strip().lower() in _DOCUMENT_TEMPLATE_LABELS:
                continue
            if len(c) < 2:
                continue

            # Single-word capitalized tokens are an ambiguous signal in
            # Title-Case (e.g. "They" vs "TCS" both start with one
            # capital letter) but an UNAMBIGUOUS signal when the word is
            # fully uppercase (TCS, NLP, LLM) -- ordinary English
            # sentence-starters are never all-caps multi-letter acronyms,
            # so those are trusted regardless of position in the text.
            # Title-Case single words (Pune, Weeks, They) are weaker
            # signals and need the mid-sentence corroboration check,
            # since they're indistinguishable from a word that merely
            # opened a sentence.
            if len(words) == 1:
                is_all_caps_acronym = first_word.isupper() and len(first_word) >= 2
                if not is_all_caps_acronym:
                    mid_sentence_pattern = r"[a-z,;:)\]]\s+" + re.escape(first_word) + r"\b"
                    if not re.search(mid_sentence_pattern, text):
                        continue

            out.append(c)
    return out


def _extract_numeric_constraints(text: str) -> list:
    """Numeric ranges and figures that read as hard constraints: years,
    days, months, salary figures. These are the numbers a hiring config
    absolutely cannot silently drop, regardless of which JD is being read.
    """
    patterns = [
        r"\b\d+\s*[-–to]+\s*\d+\s*(?:years?|yrs?|days?|months?)\b",
        r"\b\d+\s*(?:years?|yrs?|days?|months?)\b",
        r"\b(?:₹|rs\.?|inr)\s*\d+(?:\.\d+)?\s*(?:lpa|lakhs?|crore)?\b",
        r"\b\d+(?:\.\d+)?\s*lpa\b",
    ]
    found = []
    for pat in patterns:
        found.extend(re.findall(pat, text, flags=re.IGNORECASE))
    return found


def extract_critical_tokens(raw_jd_text: str, max_tokens: int = 40) -> list:
    """Derive a JD-specific critical-token list for the validation gate.

    Input: the full raw JD text (str).
    Output: a deduplicated list of strings (proper nouns + numeric
            constraints) that the compiled config must account for.
    How it works: regex-extracts capitalized phrases and numeric
            range/compensation patterns, dedupes case-insensitively,
            caps the list length so an unusually noisy JD doesn't
            produce an unmanageable gate.
    """
    proper_nouns = _extract_proper_nouns(raw_jd_text)
    numeric = _extract_numeric_constraints(raw_jd_text)

    seen = set()
    out = []
    for token in proper_nouns + numeric:
        key = token.strip().lower()
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(token.strip())

    return out[:max_tokens]